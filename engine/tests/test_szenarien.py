# -*- coding: utf-8 -*-
"""Anti-Regressions-Testmatrix: pathologische Dateien durch die echte Pipeline.

Warum es diesen Test gibt
-------------------------
Die schlimmsten Fehler dieser App sind nicht die lauten. Ein Absturz fällt auf.
Eine Datei, die still als Zeichensalat in der Bibliothek landet, fällt erst
Monate später auf – beim Zitieren. Dieser Test erzeugt deshalb **selbst** eine
Reihe schwieriger Dateien (nichts Fremdes liegt im Repo) und schickt sie durch
dieselbe Kette, die auch die App benutzt: `extract` -> `chunker` -> `normalize`
-> Datenbank -> Suche.

Für jede Zeile der Matrix gibt es ein festes Soll:

    * entweder ein brauchbares Ergebnis (Text da, Seitenzahl stimmt),
    * oder ein benannter, verständlicher Fehler (Ausnahme mit Klartext),

aber nie: Absturz, Zeichensalat oder ein halber Datenbankeintrag.

Aufruf:  PYTHONUTF8=1 python engine/tests/test_szenarien.py

Der Läufer am Ende bricht NICHT beim ersten Fehlschlag ab, sondern führt alle
Tests aus und listet die roten am Schluss auf – die Liste ist die Arbeitsliste.
"""
from __future__ import annotations

import codecs
import os
import shutil
import sys
import tempfile
import traceback
from contextlib import contextmanager
from pathlib import Path

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))

from echo_engine import connect, search                       # noqa: E402
from echo_engine import extract as extract_modul              # noqa: E402
from echo_engine.chunker import chunk_pages                   # noqa: E402
from echo_engine.indexer import (delete_documents,            # noqa: E402
                                 index_document, index_pages)
from echo_engine.normalize import normalize, to_index_forms   # noqa: E402

# Arabische/japanische Ausgaben dürfen die Windows-Konsole nicht sprengen –
# ein UnicodeEncodeError im print() sähe aus wie ein Testfehler.
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass


class Uebersprungen(Exception):
    """Voraussetzung fehlt (z. B. ein noch nicht gebautes Modul).

    Wird vom Läufer gesondert gemeldet – ein übersprungener Test ist kein
    bestandener Test, aber auch kein Fehlschlag.
    """


# ---------------------------------------------------------------------------
# Hilfsmittel
# ---------------------------------------------------------------------------

MARKER = "ZITATMARKEXYZ"       # taucht in keinem Fülltext auf
MARKER_SEITE = 3
PDF_SEITEN = 5

# Fülltext der Test-PDFs: bewusst reines ASCII. Die Sonderzeichen dieses Tests
# stecken in den Dateinamen und in den TXT-Dateien, nicht in den PDF-Schriften.
_PDF_ZEILEN = [
    "Dies ist eine ganz normale deutsche Testseite mit Fliesstext.",
    "Sie enthaelt mehrere Zeilen, damit die Absatzerkennung Material hat.",
    "Der Text ist lang genug fuer die eingebaute Salat-Erkennung.",
    "Noch eine Zeile mit weiteren Woertern zur Fuellung dieser Seite.",
]


@contextmanager
def _ordner():
    """Temporärer Arbeitsordner, der hinterher wieder verschwindet."""
    d = Path(tempfile.mkdtemp(prefix="aicp-szenarien-"))
    try:
        yield d
    finally:
        shutil.rmtree(d, ignore_errors=True)


def _fitz():
    import fitz
    return fitz


def pdf_mit_marker(pfad: Path, seiten: int = PDF_SEITEN,
                   marker: str = MARKER,
                   marker_seite: int = MARKER_SEITE) -> Path:
    """Text-PDF mit bekanntem Markertext auf bekannter Seite.

    Grundlage des Seitenzahl-Versprechens: der Marker MUSS nach dem Einlesen
    auf genau dieser Seite stehen.
    """
    fitz = _fitz()
    doc = fitz.open()
    for i in range(1, seiten + 1):
        seite = doc.new_page()
        y = 72.0
        for zeile in _PDF_ZEILEN:
            seite.insert_text((72, y), f"{zeile} (Seite {i})", fontsize=11)
            y += 18
        if i == marker_seite:
            seite.insert_text((72, y), marker, fontsize=11)
    doc.save(str(pfad))
    doc.close()
    return pfad


def pdf_ohne_seiten(pfad: Path) -> Path:
    """PDF mit null Seiten – von Hand gebaut.

    PyMuPDF weigert sich, ein leeres Dokument zu speichern ("cannot save with
    zero pages"), das Gebilde kommt aber sehr wohl in freier Wildbahn vor.
    Deshalb die Objekte samt gültiger xref-Tabelle direkt schreiben.
    """
    objekte = [b"<</Type/Catalog/Pages 2 0 R>>",
               b"<</Type/Pages/Kids[]/Count 0>>"]
    buf = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for nr, obj in enumerate(objekte, start=1):
        offsets.append(len(buf))
        buf += b"%d 0 obj\n" % nr + obj + b"\nendobj\n"
    xref = len(buf)
    buf += b"xref\n0 %d\n" % (len(objekte) + 1)
    buf += b"0000000000 65535 f \n"
    for off in offsets:
        buf += b"%010d 00000 n \n" % off
    buf += (b"trailer\n<</Size %d/Root 1 0 R>>\nstartxref\n%d\n%%%%EOF\n"
            % (len(objekte) + 1, xref))
    pfad.write_bytes(bytes(buf))
    return pfad


def pdf_abgeschnitten(pfad: Path, quelle: Path, anteil: float = 0.4) -> Path:
    """Die ersten ~40 % der Bytes eines gültigen PDFs – ein abgebrochener
    Download oder eine halb kopierte Datei sieht genau so aus."""
    roh = quelle.read_bytes()
    pfad.write_bytes(roh[:int(len(roh) * anteil)])
    return pfad


def pdf_verschluesselt(pfad: Path, quelle: Path) -> Path:
    """PDF mit reinem RECHTE-Passwort (Besitzerpasswort), leerem Benutzer-
    passwort. Solche Dateien sind ganz normal lesbar – sehr viele Verlags-PDFs
    sind so geschützt. Wer sie abweist, weist halbe Bibliotheken ab.
    """
    fitz = _fitz()
    doc = fitz.open(str(quelle))
    try:
        doc.save(str(pfad), encryption=fitz.PDF_ENCRYPT_AES_256,
                 owner_pw="geheim", user_pw="",
                 permissions=int(fitz.PDF_PERM_ACCESSIBILITY
                                 | fitz.PDF_PERM_PRINT))
    finally:
        doc.close()
    return pfad


def docx_einfach(pfad: Path, text: str = "Ein einfacher Absatz mit genug "
                                         "Text fuer den Chunker.") -> Path:
    import docx
    d = docx.Document()
    d.add_paragraph(text)
    d.add_paragraph("Ein zweiter Absatz, damit es mehr als eine Zeile ist.")
    d.save(str(pfad))
    return pfad


def docx_mit_tabelle(pfad: Path, zellwort: str = "TABELLENWORTXYZ") -> Path:
    """Word-Datei, deren Kerninhalt in einer TABELLE steht.

    In arabischen Fachbüchern stehen Überlieferungsketten, Wortlisten und
    Konkordanzen fast immer in Tabellen. Geht deren Inhalt verloren, ist das
    Buch scheinbar eingelesen und trotzdem nicht durchsuchbar.
    """
    import docx
    d = docx.Document()
    d.add_paragraph("Vor der Tabelle steht ein normaler Absatz mit Text.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = zellwort
    t.cell(0, 1).text = "Zweite Spalte mit Inhalt"
    t.cell(1, 0).text = "Dritte Zelle mit Inhalt"
    t.cell(1, 1).text = "Vierte Zelle mit Inhalt"
    d.add_paragraph("Nach der Tabelle folgt noch ein Absatz.")
    d.save(str(pfad))
    return pfad


def _text_von(res) -> str:
    """Gesamten Text eines Extraktionsergebnisses zusammenziehen."""
    return " ".join((t or "") for _, t in res.pages)


def _zeichensalat(text: str) -> bool:
    """Erkennt unlesbares Ergebnis: Ersatzzeichen (U+FFFD) sind der sichere
    Beweis einer falsch geratenen Kodierung; zusätzlich die eingebaute
    Salat-Erkennung der Extraktion für zerschossene PDF-Textebenen."""
    if not text:
        return False
    if "�" in text:
        return True
    return extract_modul._looks_like_garble(text)


@contextmanager
def _ohne_word_und_konverter():
    """Simuliert einen Rechner ohne Word und ohne LibreOffice.

    Zwei Gründe: (1) genau dieser Rechner ist der Normalfall beim Anwender,
    (2) ohne die Attrappe würde der Test auf einem Entwicklungsrechner echtes
    Word starten oder – schlimmer – über `find_soffice(auto_install=True)`
    einen mehrere hundert Megabyte großen Konverter herunterladen.
    """
    alt_word = extract_modul._word_installed
    alt_konv = extract_modul.convert_docx_to_pdf
    extract_modul._word_installed = lambda: False
    extract_modul.convert_docx_to_pdf = lambda *a, **k: None
    try:
        yield
    finally:
        extract_modul._word_installed = alt_word
        extract_modul.convert_docx_to_pdf = alt_konv


# ---------------------------------------------------------------------------
# 1. PDF: Seitenzahl-Versprechen und Verlässlichkeit
# ---------------------------------------------------------------------------

def test_pdf_marker_steht_auf_der_erwarteten_seite():
    """Das Kernversprechen: Seite im PDF = Seite in der App."""
    with _ordner() as d:
        pfad = pdf_mit_marker(d / "gut.pdf")
        res = extract_modul.extract(pfad)

        assert len(res.pages) == PDF_SEITEN, \
            f"Seitenzahl falsch: {len(res.pages)} statt {PDF_SEITEN}"
        gefunden = [no for no, t in res.pages if MARKER in (t or "")]
        assert gefunden == [MARKER_SEITE], \
            f"Marker auf Seite {gefunden} statt {MARKER_SEITE}"
        assert [no for no, _ in res.pages] == list(range(1, PDF_SEITEN + 1)), \
            "Seiten sind nicht lückenlos von 1 an durchnummeriert"
        assert res.reliability == "sicher", \
            f"reliability ist '{res.reliability}' statt 'sicher'"
        assert res.engine == "PDF", f"engine ist '{res.engine}'"
        assert not res.needs_ocr, "Text-PDF wurde faelschlich als Scan erkannt"
        assert not _zeichensalat(_text_von(res)), "PDF-Text ist Zeichensalat"
    print("OK  Text-PDF: Marker auf der erwarteten Seite, reliability 'sicher'")


def test_verschluesseltes_pdf_wird_gelesen():
    """Nur-Rechte-Verschlüsselung darf kein Hindernis sein."""
    with _ordner() as d:
        quelle = pdf_mit_marker(d / "quelle.pdf")
        pfad = pdf_verschluesselt(d / "geschuetzt.pdf", quelle)
        res = extract_modul.extract(pfad)

        assert len(res.pages) == PDF_SEITEN, \
            f"nur {len(res.pages)} Seiten gelesen"
        gefunden = [no for no, t in res.pages if MARKER in (t or "")]
        assert gefunden == [MARKER_SEITE], \
            f"Marker im geschuetzten PDF auf {gefunden} statt {MARKER_SEITE}"
        assert res.reliability == "sicher"
    print("OK  Verschluesseltes PDF (nur Rechte-Passwort) wird gelesen")


# ---------------------------------------------------------------------------
# 2. Kaputte Dateien: benannter Fehler statt Absturz oder Salat
# ---------------------------------------------------------------------------

def test_kaputte_dateien_melden_klaren_fehler():
    """Vier Sorten kaputter Dateien, ein gemeinsames Soll.

    Erlaubt ist: eine Ausnahme mit verständlicher Klartext-Meldung ODER ein
    Ergebnis ohne Seiten. Nicht erlaubt: eine Ausnahme ohne Meldung, eine
    Ausnahme, die den Prozess sprengt (SystemExit/KeyboardInterrupt), oder
    Seiten voller Zeichensalat.
    """
    with _ordner() as d:
        gut = pdf_mit_marker(d / "gut.pdf")

        leer = d / "leer.pdf"
        leer.write_bytes(b"")
        kein_pdf = d / "keinpdf.pdf"
        kein_pdf.write_text("Das ist gar kein PDF, sondern nur Text.\n" * 20,
                            encoding="utf-8")

        faelle = {
            "0-Byte-Datei": leer,
            "abgeschnittenes PDF": pdf_abgeschnitten(d / "halb.pdf", gut),
            "kein PDF trotz .pdf": kein_pdf,
            "PDF ohne Seiten": pdf_ohne_seiten(d / "null.pdf"),
        }

        probleme: list[str] = []
        for name, pfad in faelle.items():
            try:
                res = extract_modul.extract(pfad)
            except (SystemExit, KeyboardInterrupt) as e:      # nie erlaubt
                probleme.append(f"{name}: reisst den Prozess ab ({e!r})")
                continue
            except Exception as e:
                grund = str(e).strip()
                if not grund:
                    probleme.append(
                        f"{name}: Ausnahme {type(e).__name__} ohne Meldung")
                else:
                    print(f"    {name}: {type(e).__name__}: {grund[:90]}")
                continue
            text = _text_von(res)
            if _zeichensalat(text):
                probleme.append(f"{name}: Zeichensalat statt Fehler "
                                f"({text[:60]!r})")
            elif res.pages and text.strip():
                print(f"    {name}: gelesen, {len(res.pages)} Seiten")
            else:
                print(f"    {name}: 0 Seiten, kein Text (in Ordnung)")

        assert not probleme, "Kaputte Dateien ohne klares Verhalten:\n  - " \
            + "\n  - ".join(probleme)
    print("OK  Kaputte Dateien: benannter Fehler oder leeres Ergebnis")


def test_dateinamen_mit_sonderzeichen():
    """Umlaute, arabische Zeichen und Leerzeichen im Dateinamen.

    Der Anwender benennt seine Bücher auf Deutsch und Arabisch. Ein Pfad, der
    nur in ASCII funktioniert, ist hier kein Randfall, sondern der Alltag.
    """
    namen = ["Übung mit Leerzeichen.pdf",
             "كتاب عربي مع مسافات.pdf",
             "Gemischt عربي und Ümlaut.pdf"]
    with _ordner() as d:
        probleme: list[str] = []
        for name in namen:
            pfad = pdf_mit_marker(d / name)
            try:
                res = extract_modul.extract(pfad)
                gefunden = [no for no, t in res.pages if MARKER in (t or "")]
                if gefunden != [MARKER_SEITE]:
                    probleme.append(f"{name}: Marker auf {gefunden}")
            except Exception as e:
                probleme.append(f"{name}: {type(e).__name__}: {e}")

        # …und dasselbe für TXT und DOCX (andere Lesewege, gleicher Pfadstress)
        txt = d / "Notizen عربي Ü.txt"
        txt.write_text("Ein kurzer Merksatz mit " + MARKER, encoding="utf-8")
        try:
            assert MARKER in _text_von(extract_modul.extract(txt))
        except Exception as e:
            probleme.append(f"{txt.name}: {type(e).__name__}: {e}")

        dx = d / "Aufsatz عربي Ü.docx"
        docx_einfach(dx, "Ein Absatz mit dem Wort " + MARKER + " darin.")
        try:
            with _ohne_word_und_konverter():
                assert MARKER in _text_von(extract_modul.extract(dx))
        except Exception as e:
            probleme.append(f"{dx.name}: {type(e).__name__}: {e}")

        assert not probleme, "Dateinamen mit Sonderzeichen:\n  - " \
            + "\n  - ".join(probleme)
    print("OK  Dateinamen mit Umlauten, Arabisch und Leerzeichen")


# ---------------------------------------------------------------------------
# 3. TXT: Kodierungen
# ---------------------------------------------------------------------------

# Deutsche und arabische Probewörter, an denen sich eine falsch geratene
# Kodierung sofort zeigt.
_DE_PROBE = "Müller aus Köln trinkt Öl auf der Straße. "
_AR_PROBE = "مرحبا زنجبيل كتاب المكتبة "


def _txt_faelle() -> dict[str, tuple[bytes, list[str]]]:
    """Kodierung -> (Dateiinhalt, Wörter die danach lesbar sein müssen)."""
    gemischt = (_DE_PROBE * 3) + "\n" + (_AR_PROBE * 3) + "\n"
    return {
        "utf-8": (gemischt.encode("utf-8"),
                  ["Müller", "Köln", "Straße", "زنجبيل"]),
        "utf-8-bom": (codecs.BOM_UTF8 + gemischt.encode("utf-8"),
                      ["Müller", "Köln", "Straße", "زنجبيل"]),
        "utf-16-le": (codecs.BOM_UTF16_LE + gemischt.encode("utf-16-le"),
                      ["Müller", "Köln", "Straße", "زنجبيل"]),
        "cp1252": ((_DE_PROBE * 4).encode("cp1252"),
                   ["Müller", "Köln", "Öl", "Straße"]),
        "cp1256": ((_AR_PROBE * 4).encode("cp1256"),
                   ["مرحبا", "زنجبيل", "المكتبة"]),
    }


def test_txt_kodierungen_ergeben_lesbaren_text():
    """Jede der fünf gängigen Kodierungen muss korrekt gelesen werden.

    Word und Notepad speichern je nach Alter und Sprache in UTF-16 oder in
    einer Codepage; arabische Textdateien aus älteren Programmen sind fast
    immer cp1256. Wird die Kodierung nicht erkannt, landet die Datei als
    Zeichensalat in der Bibliothek – suchbar ist sie dann nie wieder, und
    aufgefallen ist es beim Einlesen nicht.
    """
    with _ordner() as d:
        probleme: list[str] = []
        for kodierung, (roh, woerter) in _txt_faelle().items():
            pfad = d / f"{kodierung}.txt"
            pfad.write_bytes(roh)
            try:
                res = extract_modul.extract(pfad)
            except Exception as e:
                probleme.append(f"{kodierung}: {type(e).__name__}: {e}")
                continue
            text = _text_von(res)
            if text.startswith("﻿"):
                probleme.append(f"{kodierung}: BOM steht im Text")
            fehlt = [w for w in woerter if w not in text]
            if fehlt:
                probleme.append(
                    f"{kodierung}: fehlende Woerter {fehlt} – gelesen wurde "
                    f"{text[:50]!r}")
            elif _zeichensalat(text):
                probleme.append(f"{kodierung}: Zeichensalat trotz Treffer")
            else:
                print(f"    {kodierung}: gelesen")

        assert not probleme, "TXT-Kodierungen:\n  - " + "\n  - ".join(probleme)
    print("OK  TXT in allen fuenf Kodierungen lesbar")


def test_txt_reliability_ist_ungefaehr():
    """Textdateien haben keine gedruckten Seiten – das muss angezeigt bleiben,
    sonst würde eine künstliche 2000-Zeichen-Seite als zitierfähig gelten."""
    with _ordner() as d:
        pfad = d / "notiz.txt"
        pfad.write_text("Ein Satz. " * 50, encoding="utf-8")
        res = extract_modul.extract(pfad)
        assert res.reliability == "ungefähr", \
            f"reliability ist '{res.reliability}' statt 'ungefähr'"
        assert res.real_page_numbers is False
    print("OK  TXT wird als 'ungefaehr' gekennzeichnet")


# ---------------------------------------------------------------------------
# 4. DOCX ohne Word und ohne Konverter (simulierte Umgebung, Punkt 5b)
# ---------------------------------------------------------------------------

def test_docx_ohne_word_und_ohne_konverter():
    """Auch ganz ohne Konverter muss eine Word-Datei Text liefern.

    Simuliert: Word nicht installiert, LibreOffice nicht vorhanden. Übrig
    bleibt der reine python-docx-Weg. Er darf keine Ausnahme werfen und muss
    den Text mitbringen – mit ehrlich als 'ungefähr' gekennzeichneten Seiten.
    """
    with _ordner() as d:
        pfad = docx_einfach(d / "einfach.docx",
                            "Ein einfacher Absatz mit dem Wort " + MARKER + ".")
        with _ohne_word_und_konverter():
            res = extract_modul.extract(pfad)

        text = _text_von(res)
        assert res.pages, "keine Seiten geliefert"
        assert MARKER in text, f"Absatztext fehlt: {text[:80]!r}"
        assert res.reliability == "ungefähr", \
            f"reliability ist '{res.reliability}' statt 'ungefähr'"
        assert res.real_page_numbers is False, \
            "geschaetzte Seiten duerfen nicht als echt gelten"
        assert res.warnings, "Notweg ohne Warnhinweis"
    print("OK  DOCX ohne Word/Konverter liefert Text (Seiten 'ungefaehr')")


def test_docx_tabelleninhalt_geht_nicht_verloren():
    """Tabellenzellen sind Inhalt, nicht Dekoration.

    Geprüft wird bewusst der Notweg ohne Word: dort sammelt die App nur
    `document.paragraphs` ein, und Tabellen stehen in Word-XML ausserhalb
    davon. Ein Buch voller Tabellen wirkt dann eingelesen und ist trotzdem
    leer.
    """
    zellwort = "TABELLENWORTXYZ"
    with _ordner() as d:
        pfad = docx_mit_tabelle(d / "tabelle.docx", zellwort)
        with _ohne_word_und_konverter():
            res = extract_modul.extract(pfad)

        text = _text_von(res)
        assert "Vor der Tabelle" in text, "nicht einmal der Absatztext kam an"
        assert zellwort in text, \
            f"Tabelleninhalt fehlt vollstaendig – gelesen: {text[:120]!r}"
        assert "Vierte Zelle" in text, "nur ein Teil der Tabelle kam an"
    print("OK  DOCX: Tabelleninhalt bleibt erhalten")


# ---------------------------------------------------------------------------
# 5. Sprachen: Tokens, Passagen, Auffindbarkeit
# ---------------------------------------------------------------------------

# Je Sprache: Probewörter + ein Fülltext, der lang genug für den Chunker ist.
SPRACHEN: dict[str, dict] = {
    "Deutsch": {
        "woerter": ["Müller", "Öl", "Straße", "Küchenschränke"],
        "text": "Die Moeglichkeiten der Kueche in Koeln sind gross. " * 8,
        "suchwort": "Küchenschränke",
    },
    "Arabisch": {
        "woerter": ["زنجبيل", "المكتبة", "كتاب"],
        "text": "هذا نص عربي طويل بما يكفي كي لا يسقط المقطع من الفهرس. " * 6,
        "suchwort": "زنجبيل",
    },
    "Persisch": {
        "woerter": ["پژوهشگر", "کتابخانه", "نوشتن"],
        "text": "این یک متن فارسی است که به اندازه کافی طولانی است برای فهرست. " * 6,
        "suchwort": "پژوهشگر",
    },
    "Kyrillisch": {
        "woerter": ["Москва", "книга", "исследование"],
        "text": "Это русский текст достаточной длины для индексации книги. " * 6,
        "suchwort": "Москва",
    },
    "CJK": {
        "woerter": ["図書館", "日本語", "研究"],
        "text": "これは十分な長さの日本語のテキストです。図書館の本について。" * 8,
        "suchwort": "図書館",
    },
}


def test_woerter_bleiben_ganze_token():
    """Ein Wort ist ein Token – auch mit Umlaut, auch ohne lateinische Schrift.

    Zerfällt "Müller" im Index in "M" und "ller", bekommt der Anwender bei der
    Suche nach "Müller" jedes Wort mit einem einzelnen "M" als Treffer, und
    "Öl" wird zu "l". Sprachen ohne arabische oder lateinische Buchstaben
    (Kyrillisch, CJK) erzeugen gar kein Token: sie sind unauffindbar.
    """
    probleme: list[str] = []
    for sprache, angaben in SPRACHEN.items():
        for wort in angaben["woerter"]:
            norm, stems = to_index_forms(wort)
            erwartet = normalize(wort)
            if norm != erwartet:
                probleme.append(
                    f"{sprache}: {wort!r} -> {norm!r} statt {erwartet!r}")
            elif not stems.strip():
                probleme.append(f"{sprache}: {wort!r} hat keinen Stamm")
    assert not probleme, "Zerbrochene Tokens:\n  - " + "\n  - ".join(probleme)
    print("OK  Woerter bleiben in allen Sprachen ganze Tokens")


def test_jede_sprache_erzeugt_passagen():
    """Ein voller Text darf nie als 'kein Text gefunden' durchfallen.

    Der Chunker verwirft Abschnitte mit weniger als 20 echten Buchstaben –
    gezählt werden nur arabische und lateinische. Ein russisches oder
    japanisches Buch hat danach null Passagen, obwohl der Text vollständig da
    ist.
    """
    probleme: list[str] = []
    for sprache, angaben in SPRACHEN.items():
        passagen = chunk_pages([(1, angaben["text"])])
        if not passagen:
            probleme.append(f"{sprache}: 0 Passagen aus "
                            f"{len(angaben['text'])} Zeichen")
    assert not probleme, "Text ohne Passagen:\n  - " + "\n  - ".join(probleme)
    print("OK  Jede Sprache erzeugt Passagen")


def test_jede_sprache_ist_wiederfindbar():
    """Ende-zu-Ende je Sprache: indexieren, suchen, genau das Buch bekommen.

    Zweiter, ebenso wichtiger Prüfpunkt: eine Suche, deren Wort die
    Zerlegung nicht kennt, ergibt eine LEERE Anfrage – und eine leere Anfrage
    blättert durch die ganze Bibliothek. Der Anwender sieht dann Treffer, die
    sein Suchwort nicht enthalten, und hält sie für Fundstellen.
    """
    con = connect(":memory:")
    for sprache, angaben in SPRACHEN.items():
        text = angaben["text"] + " " + " ".join(angaben["woerter"])
        index_pages(con, [(1, text)], title=sprache, author="Test")

    probleme: list[str] = []
    for sprache, angaben in SPRACHEN.items():
        treffer = search(con, angaben["suchwort"], limit=20)
        titel = sorted({h.title for h in treffer})
        if not treffer:
            probleme.append(f"{sprache}: {angaben['suchwort']!r} findet nichts")
        elif titel != [sprache]:
            probleme.append(
                f"{sprache}: {angaben['suchwort']!r} liefert {titel} "
                f"(fremde Buecher als Treffer)")
    con.close()
    assert not probleme, "Suche je Sprache:\n  - " + "\n  - ".join(probleme)
    print("OK  Jede Sprache ist wiederfindbar (und nur sie)")


# ---------------------------------------------------------------------------
# 6. Ende zu Ende über die Datenbank
# ---------------------------------------------------------------------------

def test_datei_wird_indexiert_und_mit_richtiger_seite_gefunden():
    """Vom PDF bis zum Suchtreffer – die Seitenzahl muss durchreichen."""
    with _ordner() as d:
        pfad = pdf_mit_marker(d / "buch.pdf")
        con = connect(":memory:")
        doc_id = index_document(con, pfad, title="Testbuch", author="Autor A")

        zeile = con.execute("SELECT title, page_count, reliability, file_type "
                            "FROM documents WHERE id=?", (doc_id,)).fetchone()
        assert zeile["page_count"] == PDF_SEITEN, \
            f"page_count {zeile['page_count']} statt {PDF_SEITEN}"
        assert zeile["reliability"] == "sicher", \
            f"reliability '{zeile['reliability']}' statt 'sicher'"
        assert zeile["file_type"] == "pdf"

        treffer = search(con, MARKER, limit=10)
        assert treffer, "eingelesenes Buch nicht auffindbar"
        assert len(treffer) == 1, f"{len(treffer)} Treffer statt einem"
        h = treffer[0]
        assert h.page_from == MARKER_SEITE, \
            f"Treffer auf Seite {h.page_from} statt {MARKER_SEITE}"
        assert h.page_from == h.page_to, "Passage laeuft ueber Seitengrenzen"
        assert h.reliability == "sicher"
        assert h.title == "Testbuch"
        con.close()
    print("OK  Datei -> Index -> Suche mit korrekter Seitenzahl")


def test_kaputte_datei_hinterlaesst_keinen_halben_eintrag():
    """Scheitert das Einlesen, darf NICHTS in der Bibliothek stehenbleiben.

    Ein Dokumenteintrag ohne Seiten wäre das Schlimmste: das Buch steht in der
    Liste, lässt sich öffnen und ist leer – ohne jeden Hinweis auf den Fehler.
    """
    with _ordner() as d:
        kaputt = d / "kaputt.pdf"
        kaputt.write_bytes(b"NICHTS DAVON IST EIN PDF" * 5)
        con = connect(":memory:")
        try:
            index_document(con, kaputt, title="Kaputt")
        except Exception:
            pass        # ein Fehler ist erlaubt – Reste sind es nicht

        for tabelle in ("documents", "pages", "passages", "passages_fts"):
            anzahl = con.execute(f"SELECT COUNT(*) FROM {tabelle}").fetchone()[0]
            assert anzahl == 0, \
                f"{anzahl} Zeile(n) in {tabelle} nach gescheitertem Einlesen"
        con.close()
    print("OK  Gescheitertes Einlesen hinterlaesst keinen halben Eintrag")


def test_loeschen_hinterlaesst_keine_geistertreffer():
    """Nach dem Löschen darf kein Wort des alten Buches mehr gefunden werden.

    Der Volltextindex ist inhaltslos und SQLite vergibt Passagen-Nummern
    erneut: ein neu eingelesenes Buch erbte sonst die Wörter des gelöschten.
    Hier bewusst mit echten Dateien nachgestellt.
    """
    with _ordner() as d:
        alt = pdf_mit_marker(d / "alt.pdf", marker=MARKER)
        neu = pdf_mit_marker(d / "neu.pdf", marker="ANDERESWORTXYZ")

        con = connect(":memory:")
        alt_id = index_document(con, alt, title="Altes Buch")
        assert len(search(con, MARKER, limit=10)) == 1

        assert delete_documents(con, [alt_id]) == 1
        assert con.execute(
            "SELECT COUNT(*) FROM passages_fts").fetchone()[0] == 0, \
            "Indexzeilen des geloeschten Buches sind stehengeblieben"

        index_document(con, neu, title="Neues Buch")
        geister = search(con, MARKER, limit=10)
        assert not geister, \
            f"Geistertreffer: {[(g.title, g.snippet[:40]) for g in geister]}"
        assert len(search(con, "ANDERESWORTXYZ", limit=10)) == 1, \
            "das neue Buch ist nicht auffindbar"
        assert con.execute("SELECT COUNT(*) FROM passages_fts").fetchone()[0] \
            == con.execute("SELECT COUNT(*) FROM passages").fetchone()[0]
        con.close()
    print("OK  Loeschen hinterlaesst keine Geistertreffer")


# ---------------------------------------------------------------------------
# 7. Simulierte Umgebung: PDF ohne PyMuPDF (Punkt 5a)
# ---------------------------------------------------------------------------

def test_pdf_ohne_pymupdf_ueber_rueckfallweg():
    """Wenn PyMuPDF nicht lädt, muss ein Rückfallweg einspringen.

    Auf Windows ist genau das schon passiert: die gebündelte DLL liess sich
    nicht laden ("DLL load failed"), und damit war jedes PDF unlesbar. Der
    Rückfallweg (`echo_engine.pdf_fallback`) muss dieselben Seitenzahlen
    liefern – sonst wäre das Zitat falsch statt nur unschön.

    Fehlt das Modul noch, wird der Test übersprungen statt rot gemeldet.
    """
    try:
        import echo_engine.pdf_fallback     # noqa: F401
    except ImportError as e:
        raise Uebersprungen(
            f"Modul echo_engine.pdf_fallback fehlt noch ({e}) – "
            "Rueckfallweg noch nicht pruefbar")

    def _kein_fitz():
        raise ImportError("PyMuPDF ist hier nicht ladbar (Attrappe)")

    with _ordner() as d:
        pfad = pdf_mit_marker(d / "ohne-fitz.pdf")
        alt = extract_modul._fitz
        extract_modul._fitz = _kein_fitz
        try:
            res = extract_modul.extract(pfad)
        finally:
            extract_modul._fitz = alt

        assert len(res.pages) == PDF_SEITEN, \
            f"Rueckfallweg liefert {len(res.pages)} statt {PDF_SEITEN} Seiten"
        gefunden = [no for no, t in res.pages if MARKER in (t or "")]
        assert gefunden == [MARKER_SEITE], \
            f"Rueckfallweg: Marker auf {gefunden} statt {MARKER_SEITE}"
        assert not _zeichensalat(_text_von(res)), \
            "Rueckfallweg liefert Zeichensalat"
    print("OK  PDF ohne PyMuPDF ueber den Rueckfallweg")


# ---------------------------------------------------------------------------
# Läufer
# ---------------------------------------------------------------------------

TESTS = [
    test_pdf_marker_steht_auf_der_erwarteten_seite,
    test_verschluesseltes_pdf_wird_gelesen,
    test_kaputte_dateien_melden_klaren_fehler,
    test_dateinamen_mit_sonderzeichen,
    test_txt_kodierungen_ergeben_lesbaren_text,
    test_txt_reliability_ist_ungefaehr,
    test_docx_ohne_word_und_ohne_konverter,
    test_docx_tabelleninhalt_geht_nicht_verloren,
    test_woerter_bleiben_ganze_token,
    test_jede_sprache_erzeugt_passagen,
    test_jede_sprache_ist_wiederfindbar,
    test_datei_wird_indexiert_und_mit_richtiger_seite_gefunden,
    test_kaputte_datei_hinterlaesst_keinen_halben_eintrag,
    test_loeschen_hinterlaesst_keine_geistertreffer,
    test_pdf_ohne_pymupdf_ueber_rueckfallweg,
]


def _lauf() -> int:
    """Führt alle Tests aus – ohne beim ersten Fehlschlag abzubrechen."""
    rot: list[tuple[str, str]] = []
    grau: list[tuple[str, str]] = []
    for t in TESTS:
        try:
            t()
        except Uebersprungen as e:
            grau.append((t.__name__, str(e)))
            print(f"UEBERSPRUNGEN  {t.__name__}: {e}")
        except AssertionError as e:
            rot.append((t.__name__, str(e) or "Zusicherung verletzt"))
            print(f"ROT  {t.__name__}")
            print("     " + (str(e) or "Zusicherung verletzt").replace(
                "\n", "\n     "))
        except Exception as e:
            rot.append((t.__name__, f"{type(e).__name__}: {e}"))
            print(f"ROT  {t.__name__}  (unerwartete Ausnahme)")
            traceback.print_exc()

    print(f"\n{len(TESTS) - len(rot) - len(grau)} von {len(TESTS)} Tests gruen"
          + (f", {len(grau)} uebersprungen" if grau else ""))
    if rot:
        print("\nOffene Punkte:")
        for name, grund in rot:
            zeilen = [z.strip() for z in grund.splitlines() if z.strip()]
            print(f"  - {name}: {zeilen[0][:160] if zeilen else ''}")
            for z in zeilen[1:6]:           # Einzelfaelle mit auflisten
                print(f"      {z[:160]}")
        return 1
    print("\nAlle Szenario-Tests bestanden.")
    return 0


if __name__ == "__main__":
    sys.exit(_lauf())
