# -*- coding: utf-8 -*-
"""Extraktion: liefert für jede Datei eine Liste (seitenzahl, text).

PDF   : seitenweise über PyMuPDF; erkennt Scans (kein Textlayer -> needs_ocr)
DOCX  : bevorzugt Konvertierung nach PDF via LibreOffice (echte Seitenzahlen);
        Fallback: reiner Text ohne verlässliche Seiten (wird markiert)
TXT   : eine künstliche "Seite" pro ~2000 Zeichen
OCR   : Tesseract (ara), falls installiert – für eingescannte PDFs
"""
from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
import tempfile
import threading
from dataclasses import dataclass, field
from pathlib import Path

# PyMuPDF (fitz) wird ERST BEI BEDARF geladen (~100 ms Importkosten). Es wird nur
# beim Einlesen/OCR gebraucht, nie beim App-Start – so erscheint das Fenster früher.
def _fitz():
    import fitz  # PyMuPDF
    return fitz

from .textlayout import (clean_text, join_wrapped_lines,
                         paragraphs_from_boxes, paragraphs_from_groups)

# Windows: Unterprozesse (Tesseract, LibreOffice) OHNE aufpoppendes
# Konsolenfenster starten. Ohne dieses Flag blitzt bei JEDEM Aufruf ein
# schwarzes Fenster auf – beim seitenweisen OCR also fortlaufend, was das
# Einlesen extrem störend macht. Auf macOS/Linux ist das Flag 0 (wirkungslos).
_NO_WINDOW = getattr(subprocess, "CREATE_NO_WINDOW", 0)


@dataclass
class ExtractResult:
    pages: list[tuple[int, str]] = field(default_factory=list)
    needs_ocr: bool = False
    used_ocr: bool = False     # wurde der Text tatsächlich per OCR erkannt?
    real_page_numbers: bool = True
    # Wie verlässlich sind die Seitenzahlen?
    #   "exakt"    – von Words eigener Engine (Word installiert oder Cloud)
    #   "ungefähr" – von LibreOffice gerendert (kann leicht abweichen)
    #   "sicher"   – direkt aus einem PDF (die gedruckten Seiten)
    reliability: str = "sicher"
    engine: str = ""
    warnings: list[str] = field(default_factory=list)


# Alle Endungen, die die App annimmt. .docm/.doc/.rtf/.odt gehen über den
# Word-/LibreOffice-Weg; ist keiner davon vorhanden, sagt die App das im
# Klartext, statt die Datei kommentarlos abzulehnen.
UNTERSTUETZTE_ENDUNGEN = (".pdf", ".docx", ".docm", ".doc", ".rtf", ".odt",
                          ".txt", ".md")


def extract(path: str | Path, force_ocr: bool = False,
            progress=None) -> ExtractResult:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(p, force_ocr=force_ocr, progress=progress)
    if suffix in (".docx", ".docm"):
        return extract_docx(p, progress=progress)
    if suffix in (".doc", ".rtf", ".odt"):
        # Ältere bzw. fremde Textformate: python-docx kann sie nicht lesen,
        # Word und LibreOffice schon. Deshalb hier ohne den python-docx-Notweg.
        return extract_ueber_konverter(p, progress=progress)
    if suffix in (".txt", ".md"):
        return extract_txt(p)
    raise ValueError(f"Nicht unterstützter Dateityp: {suffix}")


def extract_ueber_konverter(path: Path, progress=None) -> ExtractResult:
    """Wandelt ein Format, das nur Word/LibreOffice lesen können, nach PDF und
    liest es von dort. Ohne Konverter gibt es eine verständliche Meldung –
    vorher wurden solche Dateien gar nicht erst angenommen (Meldung „– ?")."""
    with tempfile.TemporaryDirectory() as tmp:
        tmpd = Path(tmp)
        if _word_installed():
            try:
                if progress:
                    progress("wird mit Word umgewandelt …")
                pdf = convert_with_word(path, tmpd)
                if pdf is not None:
                    res = extract_pdf(pdf)
                    res.reliability = "exakt"
                    res.engine = "Word"
                    return res
            except Exception:
                import traceback
                traceback.print_exc()
        try:
            pdf = convert_docx_to_pdf(path, tmpd, progress)
            if pdf is not None:
                res = extract_pdf(pdf)
                res.reliability = "ungefähr"
                res.engine = "LibreOffice"
                res.warnings.append(
                    "Mit LibreOffice gewandelt – Seitenzahlen können bei "
                    "langen Dokumenten leicht abweichen.")
                return res
        except Exception:
            import traceback
            traceback.print_exc()
    raise RuntimeError(
        f"Für {path.suffix.upper()}-Dateien wird Microsoft Word oder "
        f"LibreOffice benötigt. Bitte die Datei in Word als .docx speichern "
        f"oder LibreOffice installieren.")


def _text_layer_broken(pages: list[tuple[int, str]]) -> bool:
    """Erkennt kaputte arabische Textschichten (visuelle/verdrehte
    Reihenfolge) über mehrere unabhängige Indizien:

    1. يف häufiger als في (häufigstes Wort rückwärts)
    2. Wörter, die mit ة (Ta Marbuta) BEGINNEN – im echten Arabisch
       unmöglich, in verdrehtem Text sehr häufig
    3. Mehr Wörter, die auf لا enden, als Wörter, die mit ال beginnen
       (rückwärts gedrehter Artikel)
    """
    full = " ".join(t for _, t in pages)
    # Sprachunabhängiger Salat-Test ZUERST: eine kaputte Zeichentabelle, die
    # lateinische Glyphen ausgibt, hat gar keine arabischen Token und würde
    # sonst unten fälschlich als „nicht kaputt" durchgehen.
    if _looks_like_garble(full):
        return True
    tokens = re.findall(r"[ء-ي]+", full)
    if len(tokens) < 30:
        return False
    fi = sum(1 for t in tokens if t == "في")
    fi_rev = sum(1 for t in tokens if t == "يف")
    ta_start = sum(1 for t in tokens if t.startswith("ة"))
    al_start = sum(1 for t in tokens if t.startswith("ال"))
    la_end = sum(1 for t in tokens if t.endswith("لا") and len(t) > 3)

    indizien = 0
    if fi_rev > max(3, fi):
        indizien += 1
    if ta_start > max(4, len(tokens) * 0.003):
        indizien += 1
    if la_end > max(4, al_start):
        indizien += 1
    return indizien >= 1


def _looks_like_garble(full: str) -> bool:
    """Erkennt eine zerschossene Textebene an untypisch vielen Einzelzeichen
    und Symbolen. Bewusst konservativ, damit normaler deutscher/englischer/
    arabischer Text nie fälschlich als kaputt gilt (dort ist der Anteil an
    Einzelzeichen und Symbolen niedrig, die Durchschnittswortlänge hoch)."""
    if len(full) < 400:
        return False
    woerter = re.findall(r"[^\W\d_]+", full, flags=re.UNICODE)
    if len(woerter) < 60:
        return False
    einzel_anteil = sum(1 for w in woerter if len(w) == 1) / len(woerter)
    symbol_anteil = sum(
        1 for c in full if not c.isalnum() and not c.isspace()) / len(full)
    schnitt = sum(len(w) for w in woerter) / len(woerter)
    if einzel_anteil > 0.45 and symbol_anteil > 0.12:
        return True
    if schnitt < 1.9 and symbol_anteil > 0.10:
        return True
    return False


def _pdf_page_lines(page) -> list[tuple]:
    """Zeilen einer PDF-Seite mit Position: (text, x0, y0, x1, y1).

    Grundlage für die Absatzerkennung – ohne die Positionen wäre nicht
    erkennbar, ob ein Zeilenumbruch ein Absatzende oder nur ein Umbruch
    innerhalb des Absatzes ist.
    """
    lines: list[tuple] = []
    d = page.get_text("dict", sort=True)
    for block in d.get("blocks", []):
        if block.get("type") != 0:      # 1 = Bild
            continue
        for line in block.get("lines", []):
            text = "".join(s.get("text", "") for s in line.get("spans", []))
            if not text.strip():
                continue
            x0, y0, x1, y1 = line.get("bbox", (0.0, 0.0, 0.0, 0.0))
            lines.append((text, x0, y0, x1, y1))
    return lines


def _pdf_page_text(page) -> str:
    """Seitentext als Absätze. Fällt bei Problemen auf den Rohtext zurück.

    Auch der Rückfall ist abgesichert: eine einzelne Seite mit defektem
    Inhaltsstrom oder kaputter Schrift darf höchstens DIESE Seite kosten,
    niemals das ganze Buch (vorher riss sie die komplette Extraktion mit)."""
    try:
        text = paragraphs_from_boxes(_pdf_page_lines(page))
        if text:
            return text
    except Exception:
        import traceback
        traceback.print_exc()
    try:
        return clean_text(page.get_text("text", sort=True))
    except Exception:
        import traceback
        traceback.print_exc()
        return ""


def _extract_pdf_ersatzweg(path: Path, urspruenglich: Exception,
                           progress=None) -> ExtractResult:
    """PDF ohne die native Bibliothek lesen (siehe pdf_fallback)."""
    from . import pdf_fallback
    if not pdf_fallback.verfuegbar():
        # Auch der Ersatzweg fehlt: dann den ECHTEN Grund weiterreichen,
        # damit im Fehlerbericht steht, was auf dem Gerät nicht lädt.
        raise urspruenglich
    res = ExtractResult()
    res.reliability = "sicher"      # Seitengrenzen kommen aus dem PDF selbst
    res.engine = "PDF (Ersatzweg)"
    res.pages = pdf_fallback.seiten_lesen(path, progress=progress)
    res.warnings.append(
        "Die eingebaute PDF-Bibliothek konnte auf diesem Gerät nicht geladen "
        "werden; die Datei wurde über den Ersatzweg gelesen. Seitenzahlen "
        "stimmen, die Absatzaufteilung kann gröber sein.")
    print(f"PDF-Ersatzweg benutzt ({type(urspruenglich).__name__}: "
          f"{urspruenglich})", flush=True)
    return res


def extract_pdf(path: Path, force_ocr: bool = False,
                progress=None) -> ExtractResult:
    res = ExtractResult()
    res.reliability = "sicher"      # PDF = feste, gedruckte Seiten
    res.engine = "PDF"
    try:
        fitz = _fitz()
    except Exception as e:
        # Die native PDF-Bibliothek lässt sich auf diesem Gerät nicht laden.
        # Statt jede PDF-Datei aufzugeben, den reinen Python-Weg nehmen: die
        # Seitenzahlen bleiben exakt, nur die Absatzerkennung ist gröber.
        import traceback
        traceback.print_exc()
        return _extract_pdf_ersatzweg(path, e, progress=progress)
    with fitz.open(path) as doc:
        # Nur rechte-beschränkte PDFs (kein echtes Benutzerpasswort) lassen sich
        # mit leerem Passwort freischalten – ein häufiger Fall, der bisher als
        # „passwortgeschützt" abgewiesen wurde.
        if getattr(doc, "is_encrypted", False):
            try:
                doc.authenticate("")
            except Exception:
                pass
        gesamt = len(doc)
        empty_pages = 0
        for i, page in enumerate(doc, start=1):
            text = _pdf_page_text(page)
            if not text:
                empty_pages += 1
            res.pages.append((i, text))
            if progress:
                progress("verarbeite", i, gesamt, "verarbeite")
        is_scan = gesamt > 0 and empty_pages / gesamt > 0.5
    broken = _text_layer_broken(res.pages)
    if broken:
        res.warnings.append(
            "Kaputte Textschicht erkannt (falsche Zeichenreihenfolge).")

    if force_ocr or is_scan or broken:
        res.needs_ocr = True
        # Sprache automatisch aus dem Schriftsystem bestimmen (Textebene bzw.
        # OSD) – kein Festnageln auf eine Sprache.
        # Die GESAMTE OCR-Kette ist abgesichert: schlägt sie fehl (Zeitablauf
        # auf einer riesigen Seite, fehlende Sprachdaten, abgestürztes
        # Tesseract), behalten wir die vorhandene Textschicht. Vorher riss ein
        # einziger Zeitablauf das ganze Buch mit – obwohl brauchbarer Text da war.
        try:
            ocr_pages = _try_ocr(path, _ocr_sprache(res.pages, path),
                                 progress=progress)
        except Exception:
            import traceback
            traceback.print_exc()
            ocr_pages = None
        # Die Textschicht NUR ersetzen, wenn OCR wirklich Text geliefert hat.
        # Sonst (OCR nicht verfügbar oder ohne Ergebnis) die vorhandene – ggf.
        # verstümmelte – Textschicht behalten. NIE leere Seiten speichern, wenn
        # eine Textschicht existiert (sonst zeigt der Leser nur „—").
        if ocr_pages is not None and any((t or "").strip() for _, t in ocr_pages):
            # Sicherheitsnetz: hier läuft jedes OCR-Ergebnis noch einmal durch
            # dieselbe Zeichen-/Leerraum-Bereinigung wie die Textschicht.
            res.pages = [(no, clean_text(t)) for no, t in ocr_pages]
            res.needs_ocr = False
            res.used_ocr = True
            res.warnings.append("Text per OCR (Texterkennung) erfasst.")
        else:
            res.warnings.append(
                "OCR ohne Ergebnis – vorhandene Textschicht beibehalten.")
    return res


def _find_soffice(progress=None) -> str | None:
    """Eingebaute/selbstgeladene Komponente bevorzugen (siehe components.py)."""
    from .components import find_soffice
    return find_soffice(auto_install=True, progress=progress)


# Harte Zeitgrenze für alles, was über die Word-Automation läuft. Ohne sie
# genügt EIN passwortgeschütztes Dokument, um einen Arbeiter dauerhaft zu
# blockieren: Word zeigt dann einen unsichtbaren (Visible=False!) Dialog, auf
# den niemand klicken kann. Bei zwei solchen Dateien steht das Einlesen komplett.
WORD_TIMEOUT = 300          # 5 Minuten je Dokument


def _winword_pids() -> set:
    """Laufende WINWORD.EXE-Prozesse. Damit lassen sich nach einem Abbruch
    genau die Instanzen beenden, die WIR gestartet haben – die geöffnete
    Word-Sitzung des Nutzers bleibt unangetastet."""
    if os.name != "nt":
        return set()
    try:
        out = subprocess.run(
            ["tasklist", "/FI", "IMAGENAME eq WINWORD.EXE", "/NH", "/FO", "CSV"],
            capture_output=True, text=True, timeout=20,
            encoding="utf-8", errors="replace", creationflags=_NO_WINDOW)
        pids = set()
        for zeile in (out.stdout or "").splitlines():
            teile = [t.strip().strip('"') for t in zeile.split('","')]
            if len(teile) > 1 and teile[1].isdigit():
                pids.add(int(teile[1]))
        return pids
    except Exception:
        return set()


def _beende_pids(pids) -> None:
    """Beendet übrig gebliebene Word-Instanzen (nur unsere, siehe oben)."""
    for pid in pids:
        try:
            subprocess.run(["taskkill", "/PID", str(pid), "/F", "/T"],
                           capture_output=True, timeout=20,
                           creationflags=_NO_WINDOW)
        except Exception:
            pass


def _mit_zeitgrenze(arbeit, timeout: float):
    """Führt `arbeit()` mit harter Zeitgrenze aus. Liefert (fertig, wert).

    Ein hängender COM-Aufruf lässt sich nicht abbrechen – der Hintergrund-Thread
    bleibt notfalls stehen (Daemon, stirbt mit dem Programm). Entscheidend ist,
    dass der ARBEITER weiterläuft und die Warteschlange nicht blockiert.
    """
    ergebnis: dict = {}

    def lauf():
        try:
            ergebnis["wert"] = arbeit()
        except BaseException as e:          # auch COM-Fehler sauber übernehmen
            ergebnis["fehler"] = e

    t = threading.Thread(target=lauf, daemon=True)
    t.start()
    t.join(timeout)
    if t.is_alive():
        return False, None
    if "fehler" in ergebnis:
        raise ergebnis["fehler"]
    return True, ergebnis.get("wert")


def _word_installed() -> bool:
    if sys.platform == "darwin":
        return Path("/Applications/Microsoft Word.app").exists()
    if os.name == "nt":
        try:
            import winreg
            for key in (r"Word.Application\CLSID",):
                winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, key).Close()
            return True
        except Exception:
            return False
    return False


def _word_range_to_text(roh: str) -> str:
    """Word-Steuerzeichen ins Absatz-Format der App übersetzen: Absatzmarke \\r
    und vertikaler Umbruch \\x0b trennen Absätze; Zellen-(\\x07)/Seiten-(\\x0c)/
    sonstige Steuerzeichen raus. Jeder Absatz durch clean_text, Absätze durch
    Leerzeile getrennt (= ein Absatz pro Zeile)."""
    roh = (roh or "").replace("\x0b", "\r").replace("\x0c", "\r")
    sauber = []
    for a in roh.split("\r"):
        a = clean_text(a.replace("\x07", " "))
        if a.strip():
            sauber.append(a)
    return "\n\n".join(sauber)


def _word_text_by_page(path: Path) -> "list[tuple[int, str]] | None":
    """Liest den ECHTEN Word-Text SEITENWEISE über die Word-Automation (Windows).
    So bekommen Word-Dokumente perfekten Text mit Words eigener, exakter
    Paginierung – ohne Umweg über ein PDF und ohne OCR. None, wenn nicht
    verfügbar/fehlgeschlagen (dann greift die bisherige Kaskade)."""
    if os.name != "nt":
        return None
    try:
        import win32com.client as win32
    except Exception:
        return None
    WD_STAT_PAGES = 2       # wdStatisticPages
    WD_GOTO_PAGE = 1        # wdGoToPage
    WD_GOTO_ABSOLUTE = 1    # wdGoToAbsolute
    vorher = _winword_pids()

    def arbeit():
        # COM muss in JEDEM Thread eigens initialisiert werden – die Arbeit
        # läuft wegen der Zeitgrenze in einem Hintergrund-Thread.
        import pythoncom
        pythoncom.CoInitialize()
        word = None
        try:
            word = win32.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            # Leere Passwörter mitgeben: Word bricht dann mit einem Fehler ab,
            # statt einen unsichtbaren Dialog zu öffnen und ewig zu warten.
            doc = word.Documents.Open(str(path), ReadOnly=True,
                                      AddToRecentFiles=False,
                                      PasswordDocument="",
                                      WritePasswordDocument="",
                                      Visible=False)
            try:
                total = int(doc.ComputeStatistics(WD_STAT_PAGES)) or 1
                starts = [word.Selection.GoTo(WD_GOTO_PAGE, WD_GOTO_ABSOLUTE, p).Start
                          for p in range(1, total + 1)]
                grenzen = starts + [doc.Content.End]
                pages = [(i + 1, _word_range_to_text(doc.Range(grenzen[i], grenzen[i + 1]).Text))
                         for i in range(total)]
            finally:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            return pages if any(t.strip() for _, t in pages) else None
        finally:
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    try:
        fertig, pages = _mit_zeitgrenze(arbeit, WORD_TIMEOUT)
    except Exception:
        import traceback
        traceback.print_exc()
        _beende_pids(_winword_pids() - vorher)
        return None
    if not fertig:
        # Word hängt (Dialog, Aktivierung, defektes Dokument): unsere Instanz
        # beenden und die Kaskade weiterlaufen lassen.
        print(f"Word antwortet nicht (> {WORD_TIMEOUT}s) – Abbruch", flush=True)
        _beende_pids(_winword_pids() - vorher)
        return None
    return pages


def convert_with_word(path: Path, out_dir: Path) -> Path | None:
    """Wandelt eine Word-Datei mit WORDS EIGENER Engine nach PDF – dadurch
    exakt dieselben Seitenumbrüche wie in Word. Nur wenn Word installiert ist.

    macOS: über AppleScript (Word ist sandboxed → in den eigenen Container
    exportieren, dann herauskopieren).
    Windows: über COM-Automation, komplett unsichtbar (Visible=False).
    """
    out = out_dir / (path.stem + ".pdf")

    if os.name == "nt":
        vorher = _winword_pids()

        def arbeit():
            import pythoncom
            import win32com.client as win32
            pythoncom.CoInitialize()
            word = None
            try:
                word = win32.DispatchEx("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0
                doc = word.Documents.Open(str(path), ReadOnly=True,
                                          AddToRecentFiles=False,
                                          PasswordDocument="",
                                          WritePasswordDocument="",
                                          Visible=False)
                try:
                    doc.SaveAs(str(out), FileFormat=17)   # 17 = wdFormatPDF
                finally:
                    try:
                        doc.Close(False)
                    except Exception:
                        pass
                return out if out.exists() else None
            finally:
                if word is not None:
                    try:
                        word.Quit()
                    except Exception:
                        pass
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

        try:
            fertig, ergebnis = _mit_zeitgrenze(arbeit, WORD_TIMEOUT)
        except Exception:
            import traceback
            traceback.print_exc()
            _beende_pids(_winword_pids() - vorher)
            return None
        if not fertig:
            print(f"Word antwortet nicht (> {WORD_TIMEOUT}s) – Abbruch", flush=True)
            _beende_pids(_winword_pids() - vorher)
            return None
        return ergebnis

    if sys.platform == "darwin":
        # Word darf nur in seinen Container schreiben.
        container = (Path.home() / "Library" / "Containers"
                     / "com.microsoft.Word" / "Data" / "echoarchive-tmp")
        try:
            container.mkdir(parents=True, exist_ok=True)
        except Exception:
            return None
        tmp_pdf = container / (path.stem + ".pdf")
        tmp_pdf.unlink(missing_ok=True)
        script = (
            'with timeout of 1200 seconds\n'
            'tell application "Microsoft Word"\n'
            '  set wasRunning to running\n'
            f'  open POSIX file "{path}"\n'
            '  set theDoc to active document\n'
            f'  save as theDoc file name "{tmp_pdf}" '
            'file format format PDF\n'
            '  close theDoc saving no\n'
            'end tell\n'
            'end timeout\n')
        try:
            r = subprocess.run(["osascript", "-e", script],
                               capture_output=True, text=True, timeout=1260)
            if tmp_pdf.exists():
                import shutil as _sh
                _sh.copy(tmp_pdf, out)
                tmp_pdf.unlink(missing_ok=True)
                return out
            print("Word-Wandlung:", r.stderr.strip()[:150], flush=True)
        except Exception:
            import traceback
            traceback.print_exc()
        return None

    return None


def convert_docx_to_pdf(path: Path, out_dir: Path, progress=None) -> Path | None:
    """Wandelt eine Word-Datei in ein PDF – so wie Word es täte.

    Der Schlüssel: Der Konverter bekommt (a) die von der App mitgelieferten
    Schriften UND (b) die auf dem Rechner gefundenen Original-Schriften
    (auch aus Microsoft Offices Cloud-Font-Cache). Damit rendert er das
    Dokument mit derselben Schrift wie Word und trifft dessen Seitenumbrüche.

    Es wird NIE ein Fremdprogramm (Microsoft Word) benutzt.
    """
    soffice = _find_soffice(progress=progress)
    if not soffice:
        return None
    from .components import install_fonts_into_converter
    install_fonts_into_converter(soffice, progress)

    # Eigenes, temporäres Konverter-Profil: verhindert, dass sich der
    # Vorgang an eine bereits laufende (sichtbare) LibreOffice-Instanz
    # hängt und ein Fenster öffnet.
    # Eigenes Profil in einem festen Ordner (nicht temporär): verhindert,
    # dass sich der Vorgang an eine laufende, sichtbare LibreOffice-Instanz
    # hängt oder ein Fenster öffnet. Bewusst der normale 'soffice'-Starter
    # (NICHT soffice.bin) und nur die Standard-Headless-Flags – alles andere
    # bringt LibreOffice auf dem Mac zum Absturz.
    from .components import components_dir
    # EIGENES Profil je Vorgang: LibreOffice duldet kein gleichzeitig benutztes
    # Profil. Da zwei Dokumente parallel verarbeitet werden (MAX_WORKERS=2),
    # hängte sich der zweite Lauf an den ersten oder brach ab – ein Fehler, der
    # nur bei Mehrfach-Upload auftrat und darum schwer zu fassen war.
    prof = components_dir() / f"lo-profile-{os.getpid()}-{threading.get_ident()}"
    prof.mkdir(parents=True, exist_ok=True)
    try:
        r = subprocess.run(
            [soffice, f"-env:UserInstallation={prof.as_uri()}",
             "--headless", "--convert-to", "pdf",
             "--outdir", str(out_dir), str(path)],
            capture_output=True, timeout=1200, creationflags=_NO_WINDOW)
        # LibreOffice bildet den Ausgabenamen selbst; bei Sonderzeichen kann er
        # abweichen. Darum zuerst der erwartete Name, sonst das einzige frisch
        # entstandene PDF im Ausgabeordner.
        pdf = out_dir / (path.stem + ".pdf")
        if pdf.exists():
            return pdf
        kandidaten = sorted(out_dir.glob("*.pdf"))
        if len(kandidaten) == 1:
            return kandidaten[0]
        print("Konvertierung fehlgeschlagen:",
              (r.stderr or b"").decode("utf-8", "replace")[:200], flush=True)
        return None
    finally:
        shutil.rmtree(prof, ignore_errors=True)


def extract_docx(path: Path, progress=None) -> ExtractResult:
    """Word-Dateien – nimmt immer die genaueste verfügbare Engine:

    1. Word installiert  -> Words eigene Engine (exakte Seitenzahlen)
    2. Cloud eingerichtet -> Microsofts Word-Engine online (exakt)
    3. sonst             -> LibreOffice (läuft überall, leichte Abweichung)
    """
    def _plaintext_result() -> ExtractResult:
        """Reiner python-docx-Text – perfekt, aber Seiten nur geschätzt.
        Dient als Notlösung, wenn KEIN Word verfügbar ist ODER eine
        Konvertierung insgeheim OCR bräuchte (defekte Textebene): OCR-Text
        wäre schlechter als der saubere Word-XML-Text. DOCX wird so NIE OCR't."""
        import docx  # python-docx
        d = docx.Document(str(path))
        # Word kennt die Absätze exakt – deshalb Leerzeile statt Umbruch und
        # kein Zusammenfassen von Zeilen (das würde nur raten).
        # WICHTIG: `document.paragraphs` enthält KEINE Tabellenzellen und keine
        # Kopf-/Fußzeilen. Ein Dokument, dessen Inhalt überwiegend in Tabellen
        # steht (Listen, Verzeichnisse, Konkordanzen), landete dadurch praktisch
        # leer in der Bibliothek – ohne jede Warnung. Darum wird der Körper hier
        # in Dokumentreihenfolge abgelaufen und Tabellen werden mitgenommen.
        full = "\n\n".join(_docx_bloecke(d))
        r = _paginate_plain(full, join=False)
        r.real_page_numbers = False
        r.reliability = "ungefähr"
        return r

    with tempfile.TemporaryDirectory() as tmp:
        tmpd = Path(tmp)

        # Stufe 1: lokales Word (am genauesten, offline, kein Konto).
        if _word_installed():
            # Bevorzugt den ECHTEN Word-Text seitenweise direkt lesen – perfekter
            # Text mit Words eigener, exakter Paginierung, ohne PDF/OCR-Umweg.
            try:
                if progress:
                    progress("wird aus Word gelesen …")
                seiten = _word_text_by_page(path)
                if seiten:
                    res = ExtractResult()
                    res.pages = seiten
                    res.reliability = "exakt"
                    res.engine = "Word"
                    res.real_page_numbers = True
                    res.warnings.append(
                        "Text direkt aus Word gelesen – Seitenzahlen exakt wie in Word.")
                    return res
            except Exception:
                import traceback
                traceback.print_exc()
            # Rückfall: Word -> PDF -> Extraktion (wie bisher), falls das
            # direkte Lesen fehlschlägt.
            try:
                if progress:
                    progress("wird mit Word umgewandelt …")
                pdf = convert_with_word(path, tmpd)
                if pdf is not None:
                    res = extract_pdf(pdf)
                    if res.used_ocr:
                        # Defekte Textebene im gewandelten PDF – der saubere
                        # Word-XML-Text ist besser als jede OCR.
                        res = _plaintext_result()
                        res.warnings.append(
                            "Word-Text übernommen (PDF-Textebene defekt) – "
                            "Seiten geschätzt.")
                        return res
                    res.reliability = "exakt"
                    res.engine = "Word"
                    res.warnings.append(
                        "Mit Words eigener Engine gewandelt – "
                        "Seitenzahlen exakt wie in Word.")
                    return res
            except Exception:
                import traceback
                traceback.print_exc()

        # Stufe 2: Microsoft-Cloud (falls eingerichtet)
        try:
            from .cloud_convert import convert_via_cloud, cloud_ready
            if cloud_ready():
                if progress:
                    progress("wird online (Word-Engine) umgewandelt …")
                pdf = convert_via_cloud(path, tmpd)
                if pdf is not None:
                    res = extract_pdf(pdf)
                    if res.used_ocr:
                        res = _plaintext_result()
                        res.warnings.append(
                            "Word-Text übernommen (PDF-Textebene defekt) – "
                            "Seiten geschätzt.")
                        return res
                    res.reliability = "exakt"
                    res.engine = "Word Cloud"
                    res.warnings.append(
                        "Über Microsofts Word-Engine (online) gewandelt – "
                        "Seitenzahlen exakt wie in Word.")
                    return res
        except Exception:
            pass

        # Stufe 3: LibreOffice (immer verfügbar, leichte Abweichung möglich)
        try:
            pdf = convert_docx_to_pdf(path, tmpd, progress)
            if pdf is not None:
                res = extract_pdf(pdf)
                if res.used_ocr:
                    # LibreOffice-PDF mit defekter Textebene → OCR liefe schlecht.
                    # Sauberen Word-XML-Text nehmen, Seiten geschätzt.
                    res = _plaintext_result()
                    res.warnings.append(
                        "Word-Text übernommen (PDF-Textebene defekt) – "
                        "Seiten geschätzt.")
                    return res
                res.reliability = "ungefähr"
                res.engine = "LibreOffice"
                res.warnings.append(
                    "Mit LibreOffice gewandelt – Seitenzahlen können bei "
                    "langen Dokumenten leicht abweichen.")
                return res
        except Exception:
            import traceback
            traceback.print_exc()

    # Allerletzte Notlösung: reiner python-docx-Text, Seiten nur geschätzt.
    res = _plaintext_result()
    res.warnings.append(
        "Kein Konverter verfügbar – Text übernommen, Seiten nur geschätzt.")
    return res


def _docx_bloecke(d) -> "list[str]":
    """Alle Textblöcke einer Word-Datei in Dokumentreihenfolge – Absätze UND
    Tabellen, dazu Kopf-/Fußzeilen. python-docx' `paragraphs` allein lässt
    Tabellen und Kopfzeilen weg (siehe Aufrufer)."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def aus_koerper(eltern) -> "list[str]":
        raus: list[str] = []
        try:
            kinder = list(eltern.element.body.iterchildren())
        except AttributeError:                 # Kopf-/Fußzeilen haben keinen body
            kinder = list(eltern._element.iterchildren())
        for kind in kinder:
            marke = kind.tag.rsplit("}", 1)[-1]
            if marke == "p":
                text = Paragraph(kind, eltern).text.strip()
                if text:
                    raus.append(text)
            elif marke == "tbl":
                for zeile in Table(kind, eltern).rows:
                    # Eine Tabellenzeile als ein Absatz; Zellen durch " | "
                    # getrennt, damit die Zuordnung erhalten bleibt.
                    zellen = [z.text.strip().replace("\n", " ")
                              for z in zeile.cells]
                    # Verbundene Zellen liefert python-docx mehrfach.
                    entdoppelt: list[str] = []
                    for z in zellen:
                        if z and (not entdoppelt or entdoppelt[-1] != z):
                            entdoppelt.append(z)
                    if entdoppelt:
                        raus.append(" | ".join(entdoppelt))
        return raus

    bloecke = aus_koerper(d)
    for abschnitt in getattr(d, "sections", []):
        for bereich in ("header", "footer"):
            try:
                teil = getattr(abschnitt, bereich)
                for absatz in teil.paragraphs:
                    text = absatz.text.strip()
                    if text and text not in bloecke:
                        bloecke.append(text)
            except Exception:
                continue
    return bloecke


def text_aus_bytes(roh: bytes) -> tuple[str, str]:
    """Erkennt die Zeichenkodierung einer Textdatei. Liefert (Text, Kodierung).

    Vorher wurde stur UTF-8 mit `errors="replace"` gelesen. Folgen, alle real:
    eine mit Windows-Editor als „Unicode" (UTF-16) gespeicherte Datei wurde zu
    Zeichensalat, den die App klaglos indexierte; eine deutsche Datei in
    cp1252 verlor jeden Umlaut („Müller" -> „M?ller", damit unauffindbar);
    eine arabische Datei in Windows-1256 verlor JEDES Zeichen und galt als
    „kein Text gefunden".

    Vorgehen: erst die Bytemarke (BOM) – die ist eindeutig. Sonst der Reihe
    nach die üblichen Kodierungen streng probieren und die erste nehmen, die
    ohne Fehler durchläuft. Als letzte Rettung UTF-8 mit Ersatzzeichen.
    """
    for bom, kodierung in ((b"\xef\xbb\xbf", "utf-8-sig"),
                           (b"\xff\xfe\x00\x00", "utf-32-le"),
                           (b"\x00\x00\xfe\xff", "utf-32-be"),
                           (b"\xff\xfe", "utf-16-le"),
                           (b"\xfe\xff", "utf-16-be")):
        if roh.startswith(bom):
            try:
                return roh.decode(kodierung), kodierung
            except UnicodeDecodeError:
                break
    # UTF-16 ohne Bytemarke erkennt man zuverlässig an den vielen Nullbytes,
    # die in echtem UTF-8/Latin-Text nie vorkommen.
    probe = roh[:4096]
    if probe and probe.count(0) > len(probe) // 4:
        gerade = sum(1 for i, b in enumerate(probe) if b == 0 and i % 2 == 1)
        ungerade = probe.count(0) - gerade
        for kodierung in (("utf-16-le", "utf-16-be") if gerade >= ungerade
                          else ("utf-16-be", "utf-16-le")):
            try:
                return roh.decode(kodierung), kodierung
            except UnicodeDecodeError:
                continue
    # UTF-8 prüft sich selbst: läuft es fehlerfrei durch, ist es praktisch immer
    # richtig.
    try:
        return roh.decode("utf-8"), "utf-8"
    except UnicodeDecodeError:
        pass
    # Die Windows-Kodierungen scheitern NIE (jedes Byte ergibt ein Zeichen).
    # Reines Durchprobieren würde daher immer die erste nehmen und z. B.
    # deutschen Text arabisch lesen („Straße" -> „Straكe"). Deshalb werden alle
    # Kandidaten bewertet und der plausibelste gewinnt.
    beste, bester_wert = None, None
    for kodierung in ("cp1252", "cp1256", "cp1251", "latin-1"):
        try:
            versuch = roh.decode(kodierung)
        except UnicodeDecodeError:
            continue
        wert = _kodierung_bewerten(versuch)
        if bester_wert is None or wert > bester_wert:
            beste, bester_wert = (versuch, kodierung), wert
    if beste:
        return beste
    return roh.decode("utf-8", errors="replace"), "utf-8 (mit Ersatzzeichen)"


def _schriftart(zeichen: str) -> str:
    """Grobe Schriftsystem-Zuordnung eines Buchstabens (für die Bewertung)."""
    o = ord(zeichen)
    if o < 0x250:
        return "lateinisch"
    if 0x400 <= o <= 0x52F:
        return "kyrillisch"
    if 0x590 <= o <= 0x6FF or 0x750 <= o <= 0x77F or 0xFB50 <= o <= 0xFEFF:
        return "arabisch"
    if o >= 0x4E00:
        return "cjk"
    return "sonstige"


# In echtem deutschem/französischem Text übliche Sonderbuchstaben. Alles andere
# aus dem oberen Latin-1-Bereich ist ein starkes Zeichen dafür, dass hier in
# Wahrheit eine andere Kodierung vorliegt (z. B. arabischer Text als cp1252).
_UEBLICHE_SONDER = set("äöüÄÖÜßéèêëáàâíìîóòôúùûñçÉÈÊÁÀÂÍÓÚÑÇ«»°–—…")


def _kodierung_bewerten(text: str) -> float:
    """Wie plausibel ist dieser entschlüsselte Text? Höher ist besser."""
    if not text:
        return -1e9
    buchstaben = [c for c in text if c.isalpha()]
    if not buchstaben:
        return -1e9
    punkte = 0.01 * len(buchstaben)
    # Wörter, die Schriftsysteme mischen, entstehen fast nur durch eine falsche
    # Kodierung („Straكe") – das wiegt schwer.
    for wort in re.findall(r"[^\W\d_]+", text):
        if len({_schriftart(c) for c in wort}) > 1:
            punkte -= 5
    punkte -= 3 * sum(1 for c in text
                      if 0xC0 <= ord(c) <= 0xFF and c not in _UEBLICHE_SONDER)
    punkte -= 10 * sum(1 for c in text
                       if ord(c) < 32 and c not in "\r\n\t")
    return punkte


def extract_txt(path: Path) -> ExtractResult:
    text, kodierung = text_aus_bytes(Path(path).read_bytes())
    res = _paginate_plain(text)
    res.real_page_numbers = False
    # Textdateien haben KEINE gedruckten Seiten – die „Seiten" sind künstliche
    # 2000-Zeichen-Blöcke. Darum ausdrücklich „ungefähr", damit die Oberfläche
    # sie nicht als zitierfähige Druckseiten ausweist (sonst „exakt").
    res.reliability = "ungefähr"
    res.engine = f"Text ({kodierung})"
    res.warnings.append(
        "Textdatei ohne Druckseiten – Seitenangaben sind nur Näherungen.")
    return res


def _paginate_plain(text: str, chars_per_page: int = 2000,
                    join: bool = True) -> ExtractResult:
    """Künstliche Seiten aus reinem Text.

    Geschnitten wird weiterhin am Rohtext, damit sich die Seitengrenzen
    gegenüber früher nicht verschieben; erst danach wird der Text jeder
    Seite aufbereitet.
    """
    res = ExtractResult()
    pos, page_no = 0, 1
    while pos < len(text):
        seite = text[pos:pos + chars_per_page]
        res.pages.append(
            (page_no, join_wrapped_lines(seite) if join else clean_text(seite)))
        pos += chars_per_page
        page_no += 1
    if not res.pages:
        res.pages = [(1, "")]
    return res


def _tesseract_cmd() -> str | None:
    """Findet Tesseract: zuerst die mitgelieferte Kopie (Windows-Installer),
    dann eine systemweite Installation."""
    import sys
    if getattr(sys, "frozen", False):
        base = Path(getattr(sys, "_MEIPASS", ""))
        cand = base / "tesseract" / ("tesseract.exe" if os.name == "nt"
                                     else "tesseract")
        if cand.exists():
            return str(cand)
    return shutil.which("tesseract")


def _ocr_sprache(pages, pdf_path) -> str:
    """Bestimmt automatisch die OCR-Sprache eines Dokuments. Grundlage ist das
    Schriftsystem der (ggf. verstümmelten) Textebene – die Zeichen stimmen auch
    bei falscher Reihenfolge. Ohne brauchbare Textebene (echter Scan) entscheidet
    Tesseracts OSD anhand der ersten Seite. Einzelschrift-Dokumente bekommen NUR
    ihre Sprache (keine Kreuz-Artefakte): arabisch → 'ara', lateinisch →
    'deu+eng'. Nur wenn wirklich beide Schriften vorkommen → 'ara+deu+eng'."""
    txt = " ".join(t or "" for _, t in pages)
    ar = len(re.findall(r"[؀-ۿ]", txt))
    la = len(re.findall(r"[A-Za-z]", txt))
    if ar + la >= 20:
        anteil = ar / (ar + la)
        if anteil >= 0.85:
            return "ara"
        if anteil <= 0.15:
            return "deu+eng"
        return "ara+deu+eng"
    return _osd_sprache(pdf_path)


def _osd_sprache(pdf_path) -> str:
    """Schriftsystem der ersten Seite via Tesseract-OSD; Rückfall 'ara'."""
    cmd = _tesseract_cmd()
    if not cmd:
        return "ara"
    tessdata = Path(cmd).parent / "tessdata"
    tdata = ["--tessdata-dir", str(tessdata)] if tessdata.exists() else []
    tmp = Path(tempfile.mkdtemp(prefix="aicp-osd-"))
    try:
        with _fitz().open(pdf_path) as doc:
            if len(doc) == 0:
                return "ara"
            png = tmp / "osd.png"
            doc[0].get_pixmap(dpi=300).save(str(png))
        out = subprocess.run([cmd, *tdata, "--psm", "0", str(png), "stdout"],
                             capture_output=True, text=True,
                             encoding="utf-8", errors="replace", timeout=60,
                             creationflags=_NO_WINDOW)
        m = re.search(r"Script:\s*(\w+)", out.stdout or "")
        skript = (m.group(1) if m else "").lower()
        if skript == "arabic":
            return "ara"
        if skript == "latin":
            return "deu+eng"
        return "ara"
    except Exception:
        return "ara"
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _try_ocr(pdf_path: Path, lang: str = "ara",
             progress=None) -> list[tuple[int, str]] | None:
    """OCR-Kette: Apple Vision (macOS, eingebaut) -> Tesseract -> None.
    `lang` ist die (automatisch erkannte) Sprache für Tesseract; Vision erkennt
    die Sprache selbst und ignoriert den Parameter. `progress` meldet – wo
    vorhanden – die OCR-Seite (i/N) für den Fortschrittsbalken."""
    import sys
    if sys.platform == "darwin":
        try:
            return _ocr_pdf_vision(pdf_path, progress=progress)
        except Exception:
            import traceback
            traceback.print_exc()
    if _tesseract_cmd():
        return _ocr_pdf_tesseract(pdf_path, lang, progress=progress)
    return None


def _ocr_pdf_vision(pdf_path: Path, progress=None) -> list[tuple[int, str]]:
    """Arabische Texterkennung über das in macOS eingebaute Vision-Framework.

    Braucht: pip install pyobjc-framework-Vision (steht in requirements.txt).
    """
    import Vision
    from Foundation import NSData

    pages: list[tuple[int, str]] = []
    with _fitz().open(pdf_path) as doc:
        gesamt = len(doc)
        for i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(dpi=200)
            png = pix.tobytes("png")
            data = NSData.dataWithBytes_length_(png, len(png))
            handler = (Vision.VNImageRequestHandler.alloc()
                       .initWithData_options_(data, None))
            req = Vision.VNRecognizeTextRequest.alloc().init()
            req.setRecognitionLevel_(0)  # 0 = accurate
            req.setUsesLanguageCorrection_(True)
            try:
                req.setRecognitionLanguages_(["ar-SA", "de-DE", "en-US"])
            except Exception:
                pass
            handler.performRequests_error_([req], None)
            obs = list(req.results() or [])
            # Vision liefert normierte Boxen mit Ursprung UNTEN links. Für die
            # Absatzerkennung werden sie auf Bildpunkte mit Ursprung OBEN
            # links umgerechnet – erst dann sind waagerechte und senkrechte
            # Abstände miteinander vergleichbar.
            lines: list[tuple] = []
            for o in obs:
                cands = o.topCandidates_(1)
                if not (cands and len(cands)):
                    continue
                box = o.boundingBox()
                bx, by = box.origin.x, box.origin.y
                bw, bh = box.size.width, box.size.height
                lines.append((str(cands[0].string()),
                              bx * pix.width, (1.0 - by - bh) * pix.height,
                              (bx + bw) * pix.width, (1.0 - by) * pix.height))
            try:
                pages.append((i, paragraphs_from_boxes(lines)))
            except Exception as e:
                # Eine misslungene Seite kostet nur diese Seite (siehe Tesseract).
                pages.append((i, ""))
                print(f"OCR Seite {i} übersprungen: "
                      f"{type(e).__name__}: {e}", flush=True)
            print(f"OCR Seite {i}/{gesamt}", flush=True)
            if progress:
                progress("ocr", i, gesamt, "ocr")
    return pages


def _tesseract_tsv_to_text(tsv: str) -> str:
    """Baut aus Tesseracts TSV-Ausgabe Absätze.

    Tesseract erkennt die Absätze selbst und gibt sie in den Spalten
    block_num/par_num aus – genauer als jede Heuristik auf dem fertigen
    Text. Spalten: level, page, block, par, line, word, left, top, width,
    height, conf, text.
    """
    if not tsv:
        return ""
    gruppen: dict[tuple[int, int], dict[int, list[str]]] = {}
    reihenfolge: list[tuple[int, int]] = []
    for row in tsv.splitlines()[1:]:
        f = row.split("\t")
        if len(f) < 12 or f[0] != "5":       # 5 = Ebene "Wort"
            continue
        try:
            conf = float(f[10])
            block, par, line = int(f[2]), int(f[3]), int(f[4])
        except ValueError:
            continue
        wort = f[11].strip()
        if not wort or conf < 0:
            continue
        key = (block, par)
        if key not in gruppen:
            gruppen[key] = {}
            reihenfolge.append(key)
        gruppen[key].setdefault(line, []).append(wort)
    absaetze = [[" ".join(gruppen[k][ln]) for ln in sorted(gruppen[k])]
                for k in reihenfolge]
    return paragraphs_from_groups(absaetze)


# Obergrenze für das Seitenbild der Texterkennung. 400 dpi sind für normale
# Buchseiten ideal, sprengen aber bei sehr großen Seiten (Karten, A0-Tafeln)
# den Speicher: eine A0-Seite ergäbe bei 400 dpi rund 700 MB als Bild. Darum
# wird die Auflösung für solche Seiten so weit gesenkt, dass das Bild unter
# dieser Grenze bleibt – lieber etwas gröber erkennen als abstürzen.
OCR_MAX_PIXEL = 40_000_000      # ca. 40 Megapixel
OCR_DPI = 400


def _ocr_dpi(page) -> int:
    """Passende Auflösung für diese Seite (siehe OCR_MAX_PIXEL)."""
    try:
        r = page.rect
        breite_zoll = max(r.width, 1) / 72.0
        hoehe_zoll = max(r.height, 1) / 72.0
        pixel = breite_zoll * hoehe_zoll * OCR_DPI * OCR_DPI
        if pixel <= OCR_MAX_PIXEL:
            return OCR_DPI
        faktor = (OCR_MAX_PIXEL / pixel) ** 0.5
        return max(120, int(OCR_DPI * faktor))
    except Exception:
        return OCR_DPI


def _ocr_pdf_tesseract(pdf_path: Path, lang: str = "ara",
                       progress=None) -> "list[tuple[int, str]] | None":
    import os as _os
    cmd = _tesseract_cmd()
    env = dict(_os.environ)
    # Mitgelieferte Sprachdaten liegen im Unterordner "tessdata" neben der
    # Binärdatei. WICHTIG: Tesseract 5 will TESSDATA_PREFIX bzw. --tessdata-dir
    # DIREKT auf diesen Ordner (nicht auf den Elternordner wie noch bei v4) –
    # sonst findet es die *.traineddata nicht und liefert leeren Text
    # ("Error opening data file … ara.traineddata"). --tessdata-dir ist
    # versionsunabhängig und eindeutig.
    tessdata = Path(cmd).parent / "tessdata"
    tdata_args: list[str] = []
    if tessdata.exists():
        env["TESSDATA_PREFIX"] = str(tessdata)
        tdata_args = ["--tessdata-dir", str(tessdata)]
    # Sprache kommt vom Aufrufer (automatisch erkannt). Wichtig: KEINE
    # Schriftsysteme mischen, wo es nicht nötig ist – deu/eng auf rein arabischen
    # Seiten erzeugen massenhaft lateinische Fehl-Lesungen und verschlechtern die
    # arabische Erkennung (gemessen). Darum ara-Dokumente nur mit "ara".
    sprache = ["-l", lang]
    pages: list[tuple[int, str]] = []
    hat_text = False       # mindestens eine Seite mit echtem Text erkannt?
    # Eigener Temp-Ordner: die Seiten-PNG wird unter einem FRISCHEN Pfad
    # geschrieben, den Python NICHT offen hält. Sonst sperrt Windows die Datei,
    # und PyMuPDFs pix.save() bricht mit „cannot remove file … Permission denied"
    # ab (ein Virenscanner verschärft das). Aufräumen ist hier nie fatal.
    tmpdir = Path(tempfile.mkdtemp(prefix="aicp-ocr-"))
    try:
        with _fitz().open(pdf_path) as doc:
            gesamt = len(doc)
            for i, page in enumerate(doc, start=1):
                png = tmpdir / f"seite-{i}.png"
                # JEDE Seite einzeln absichern: eine überlange, riesige oder
                # defekte Seite kostet nur diese Seite. Vorher riss ein
                # Zeitablauf die gesamte OCR (und damit das Buch) mit.
                try:
                    pix = page.get_pixmap(dpi=_ocr_dpi(page))
                    pix.save(str(png))      # frischer Pfad, kein offenes Handle
                    # --psm 6 (ein Textblock) erkennt gemessen deutlich mehr als
                    # die Automatik (psm 3). Die TSV-Ausgabe verträgt sich NICHT
                    # mit --psm 6 (Tesseract liefert dann Plain-Text statt
                    # Spalten), daher reiner Text + zeilenbasierte Absätze.
                    out = subprocess.run(
                        [cmd, *tdata_args, str(png), "stdout", *sprache,
                         "--psm", "6"],
                        capture_output=True, text=True,
                        encoding="utf-8", errors="replace", timeout=180, env=env,
                        creationflags=_NO_WINDOW)
                    if out.returncode == 0:
                        text = join_wrapped_lines(out.stdout)
                    else:
                        # Tesseracts eigene Meldung protokollieren – sie nennt
                        # den Grund (fehlende Sprachdaten, Pfadproblem) und war
                        # bisher komplett verworfen, also von außen unsichtbar.
                        text = ""
                        fehler = (out.stderr or "").strip().splitlines()
                        print(f"OCR Seite {i}: Tesseract-Fehler "
                              f"(Code {out.returncode}) "
                              f"{fehler[0] if fehler else ''}", flush=True)
                except Exception as e:
                    text = ""
                    print(f"OCR Seite {i} übersprungen: "
                          f"{type(e).__name__}: {e}", flush=True)
                pages.append((i, text))
                if text.strip():
                    hat_text = True
                try:
                    png.unlink()            # sofort aufräumen …
                except OSError:
                    pass                    # … aber nie den Upload abbrechen
                print(f"OCR Seite {i}/{gesamt}", flush=True)
                if progress:
                    progress("ocr", i, gesamt, "ocr")
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)
    # Hat OCR NICHTS Brauchbares geliefert (z. B. Tesseract-Fehler), als
    # Fehlschlag melden: der Aufrufer behält dann die vorhandene Textebene,
    # statt leere Seiten zu speichern.
    if not hat_text:
        return None
    return pages
